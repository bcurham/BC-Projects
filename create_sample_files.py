"""
Script to create sample template and URS documents for testing
Run this after installing requirements.txt
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def create_test_script_template():
    """Create a sample test script template"""
    doc = Document()

    # Add title
    title = doc.add_heading('Test Script Template', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add some metadata
    doc.add_paragraph('Project Name: _______________________')
    doc.add_paragraph('Test Script ID: _______________________')
    doc.add_paragraph('Version: _______________________')
    doc.add_paragraph('Date: _______________________')
    doc.add_paragraph()

    # Add instructions
    doc.add_heading('Instructions', level=2)
    doc.add_paragraph(
        'This test script should be completed during software validation testing. '
        'For each test step, execute the described action and verify the expected result. '
        'Mark Pass/Fail, add your initials, and the date of execution.'
    )
    doc.add_paragraph()

    # Create table with proper headers
    doc.add_heading('Test Steps', level=2)

    # Create table (1 header row + 1 empty row for the template)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'

    # Set header row
    header_cells = table.rows[0].cells
    headers = ['Step', 'Requirement #', 'Description', 'Expected Result', 'Pass/Fail', 'Initial', 'Date']

    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Make header bold
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add approval section
    doc.add_paragraph()
    doc.add_heading('Approval', level=2)
    doc.add_paragraph('Tester: _______________________ Date: _____________')
    doc.add_paragraph('Reviewer: _______________________ Date: _____________')
    doc.add_paragraph('QA Approver: _______________________ Date: _____________')

    # Save document
    output_path = 'sample_test_script_template.docx'
    doc.save(output_path)
    print(f"✓ Created sample template: {output_path}")
    return output_path


def create_sample_urs():
    """Create a sample URS document for testing"""
    doc = Document()

    # Add title
    title = doc.add_heading('User Requirements Specification (URS)', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata
    doc.add_paragraph('Document ID: URS-001')
    doc.add_paragraph('System Name: Sample Laboratory Information System')
    doc.add_paragraph('Version: 1.0')
    doc.add_paragraph('Date: 2025-01-01')
    doc.add_paragraph()

    # Add introduction
    doc.add_heading('1. Introduction', level=2)
    doc.add_paragraph(
        'This User Requirements Specification defines the functional and non-functional '
        'requirements for the Sample Laboratory Information System.'
    )
    doc.add_paragraph()

    # Add requirements
    doc.add_heading('2. Functional Requirements', level=2)

    requirements = [
        {
            'id': 'REQ-001',
            'desc': 'User Authentication',
            'detail': 'The system shall require users to log in with a unique username and password.',
            'acceptance': 'User successfully logs in with valid credentials and is denied access with invalid credentials.'
        },
        {
            'id': 'REQ-002',
            'desc': 'Password Requirements',
            'detail': 'The system shall enforce password complexity requirements: minimum 8 characters, including uppercase, lowercase, numbers, and special characters.',
            'acceptance': 'System rejects passwords that do not meet complexity requirements.'
        },
        {
            'id': 'REQ-003',
            'desc': 'User Roles and Permissions',
            'detail': 'The system shall support multiple user roles (Administrator, Analyst, Reviewer) with different permission levels.',
            'acceptance': 'Users can only access features and data appropriate to their assigned role.'
        },
        {
            'id': 'REQ-004',
            'desc': 'Sample Registration',
            'detail': 'The system shall allow authorized users to register new samples with unique sample IDs, sample type, and collection date.',
            'acceptance': 'New samples are successfully registered with all required information and assigned unique IDs.'
        },
        {
            'id': 'REQ-005',
            'desc': 'Test Assignment',
            'detail': 'The system shall allow users to assign laboratory tests to registered samples.',
            'acceptance': 'Tests are successfully assigned to samples and appear in the work queue.'
        },
        {
            'id': 'REQ-006',
            'desc': 'Result Entry',
            'detail': 'The system shall allow analysts to enter test results with units of measurement.',
            'acceptance': 'Test results are entered successfully and stored in the database.'
        },
        {
            'id': 'REQ-007',
            'desc': 'Result Review',
            'detail': 'The system shall require results to be reviewed and approved by a second user before being finalized.',
            'acceptance': 'Results cannot be finalized without review approval from an authorized reviewer.'
        },
        {
            'id': 'REQ-008',
            'desc': 'Audit Trail',
            'detail': 'The system shall maintain an audit trail of all data entries, modifications, and deletions including user ID, date/time, and reason for change.',
            'acceptance': 'All system actions are logged with complete audit information.'
        },
        {
            'id': 'REQ-009',
            'desc': 'Data Backup',
            'detail': 'The system shall perform automated daily backups of all data.',
            'acceptance': 'Backups are created successfully each day and can be restored if needed.'
        },
        {
            'id': 'REQ-010',
            'desc': 'Report Generation',
            'detail': 'The system shall allow users to generate PDF reports of sample test results.',
            'acceptance': 'Reports are generated successfully in PDF format with all relevant sample and result information.'
        }
    ]

    for req in requirements:
        doc.add_heading(f"2.{requirements.index(req) + 1} {req['id']}: {req['desc']}", level=3)
        doc.add_paragraph(f"Requirement: {req['detail']}")
        doc.add_paragraph(f"Acceptance Criteria: {req['acceptance']}")
        doc.add_paragraph()

    # Add non-functional requirements
    doc.add_heading('3. Non-Functional Requirements', level=2)

    nf_requirements = [
        {
            'id': 'REQ-NF-001',
            'desc': 'Performance',
            'detail': 'The system shall respond to user actions within 2 seconds under normal load.',
            'acceptance': 'Response time measurements show 95% of actions complete within 2 seconds.'
        },
        {
            'id': 'REQ-NF-002',
            'desc': 'Availability',
            'detail': 'The system shall be available 99.5% of the time during business hours.',
            'acceptance': 'System uptime logs demonstrate 99.5% or greater availability.'
        },
        {
            'id': 'REQ-NF-003',
            'desc': 'Data Integrity',
            'detail': 'The system shall prevent unauthorized modification or deletion of finalized data.',
            'acceptance': 'Attempts to modify finalized data are prevented and logged.'
        }
    ]

    for req in nf_requirements:
        doc.add_heading(f"3.{nf_requirements.index(req) + 1} {req['id']}: {req['desc']}", level=3)
        doc.add_paragraph(f"Requirement: {req['detail']}")
        doc.add_paragraph(f"Acceptance Criteria: {req['acceptance']}")
        doc.add_paragraph()

    # Add compliance section
    doc.add_heading('4. Compliance Requirements', level=2)
    doc.add_paragraph(
        'This system must comply with FDA 21 CFR Part 11 for electronic records and electronic signatures.'
    )

    # Save document
    output_path = 'sample_urs_document.docx'
    doc.save(output_path)
    print(f"✓ Created sample URS: {output_path}")
    return output_path


if __name__ == '__main__':
    print("Creating sample files for testing...")
    print()

    try:
        template_path = create_test_script_template()
        urs_path = create_sample_urs()

        print()
        print("=" * 60)
        print("Sample files created successfully!")
        print("=" * 60)
        print()
        print("You can now test the application:")
        print(f"1. Upload '{urs_path}' as the URS document")
        print(f"2. Upload '{template_path}' as the template")
        print("3. Click 'Generate Test Script'")
        print()
        print("The system will extract the 13 requirements from the URS")
        print("and populate them into your test script template.")

    except Exception as e:
        print(f"Error creating sample files: {e}")
        import traceback
        traceback.print_exc()
