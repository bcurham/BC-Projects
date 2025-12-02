import os
import json
import uuid
from flask import Flask, request, jsonify, send_file, render_template, session, redirect, url_for, flash
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import anthropic
from docx import Document
import PyPDF2
from io import BytesIO
import traceback
from datetime import datetime

# Database and authentication imports
from flask_login import LoginManager, login_required, current_user
from models import db, User, Project, Template, ProjectVersion, init_db
from auth import auth_bp

# Try to import pdfplumber (optional, may fail on some systems)
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
    print("✓ pdfplumber is available for enhanced PDF extraction")
except ImportError as e:
    PDFPLUMBER_AVAILABLE = False
    print(f"⚠ pdfplumber not available (will use PyPDF2 only): {e}")
    print("  This is normal on some Windows systems. PDF extraction will still work.")

# Import new feature modules (optional enhancements)
try:
    from modules import (
        generate_rtm_excel, generate_rtm_word,
        ChangeAnalyzer, ValidationDocGenerator,
        QualityChecker, AuditPackageExporter
    )
    ENHANCED_FEATURES_AVAILABLE = True
    print("✓ Enhanced validation features loaded successfully")
except ImportError as e:
    ENHANCED_FEATURES_AVAILABLE = False
    print(f"⚠ Enhanced features not available: {e}")
    print("  Core test script generation will still work.")

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('FLASK_SECRET_KEY', 'dev-secret-key-change-in-production')
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Database configuration
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'sqlite:///testscript_generator.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize database
init_db(app)

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'auth.login'
login_manager.login_message = 'Please log in to access this page.'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Register authentication blueprint
app.register_blueprint(auth_bp)

# Initialize Anthropic client
api_key = os.getenv('ANTHROPIC_API_KEY')
if not api_key:
    print("WARNING: ANTHROPIC_API_KEY not set in environment variables!")
client = anthropic.Anthropic(api_key=api_key) if api_key else None

# Initialize enhanced feature modules (if available)
if ENHANCED_FEATURES_AVAILABLE:
    os.makedirs('baselines', exist_ok=True)  # For change analysis
    change_analyzer = ChangeAnalyzer()
    validation_doc_gen = ValidationDocGenerator()
    quality_checker = QualityChecker(client)
    audit_exporter = AuditPackageExporter()
    print("✓ Enhanced feature instances initialized")

ALLOWED_URS_EXTENSIONS = {'pdf', 'docx', 'txt'}
ALLOWED_TEMPLATE_EXTENSIONS = {'docx'}


def allowed_file(filename, allowed_extensions):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions


def extract_text_from_pdf(file_path):
    """Extract text from PDF file using PyPDF2 (with optional pdfplumber enhancement)"""
    text = ""

    # Try pdfplumber first if available (better for complex PDFs)
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(file_path) as pdf:
                print(f"Using pdfplumber: PDF has {len(pdf.pages)} pages")
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        print(f"Page {i+1}: extracted {len(page_text)} characters")

                if text.strip():
                    return text
                else:
                    print("pdfplumber returned no text, falling back to PyPDF2")
        except Exception as e:
            print(f"pdfplumber failed: {e}, falling back to PyPDF2")

    # Use PyPDF2 (works on all systems)
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            print(f"Using PyPDF2: PDF has {len(pdf_reader.pages)} pages")
            for i, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                    print(f"Page {i+1}: extracted {len(page_text)} characters")
    except Exception as e:
        print(f"PyPDF2 failed: {e}")
        raise ValueError(f"Unable to extract text from PDF. The file may be password-protected, corrupted, or contains only images. Error: {str(e)}")

    if not text.strip():
        raise ValueError("PDF file appears to be empty or contains only images. Please ensure the PDF has extractable text.")

    return text


def extract_text_from_docx(file_path):
    """Extract text from Word document including paragraphs and tables"""
    try:
        doc = Document(file_path)
        text = ""

        # Extract from paragraphs
        paragraph_count = 0
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
                paragraph_count += 1

        print(f"Extracted {paragraph_count} paragraphs from Word document")

        # Extract from tables
        table_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
            table_count += 1

        print(f"Extracted text from {table_count} tables in Word document")

        if not text.strip():
            raise ValueError("Word document appears to be empty or contains no extractable text.")

        return text
    except Exception as e:
        print(f"Error extracting from Word document: {e}")
        raise ValueError(f"Unable to read Word document. The file may be corrupted or in an unsupported format. Error: {str(e)}")


def extract_text_from_txt(file_path):
    """Extract text from plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()

        if not text.strip():
            raise ValueError("Text file is empty.")

        print(f"Extracted {len(text)} characters from text file")
        return text
    except UnicodeDecodeError:
        # Try different encodings
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            print(f"Extracted {len(text)} characters from text file using latin-1 encoding")
            return text
        except Exception as e:
            raise ValueError(f"Unable to read text file. The file encoding may not be supported. Error: {str(e)}")
    except Exception as e:
        raise ValueError(f"Unable to read text file: {str(e)}")


def extract_urs_text(file_path, filename):
    """Extract text from URS document based on file type"""
    print(f"Extracting text from: {file_path}")
    print(f"File exists: {os.path.exists(file_path)}")
    print(f"File size: {os.path.getsize(file_path) if os.path.exists(file_path) else 'N/A'} bytes")

    ext = filename.rsplit('.', 1)[1].lower()

    if ext == 'pdf':
        text = extract_text_from_pdf(file_path)
    elif ext == 'docx':
        text = extract_text_from_docx(file_path)
    elif ext == 'txt':
        text = extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    print(f"Extracted text length: {len(text)} characters")
    print(f"First 200 chars: {text[:200] if text else 'EMPTY'}")

    return text


def load_master_prompt():
    """Load the master prompt from master_prompt.md"""
    try:
        with open('master_prompt.md', 'r', encoding='utf-8') as file:
            return file.read()
    except FileNotFoundError:
        # Fallback to embedded prompt if file doesn't exist
        return """You are an expert in software validation for pharmaceutical and medical device industries.
Parse the URS document and extract all individual requirements.
Generate a detailed test script for each requirement in valid JSON format only.

Output structure:
{
  "test_steps": [
    {
      "step_no": 1,
      "requirement_id": "REQ-001",
      "description": "Requirement description",
      "expected_result": "Expected outcome"
    }
  ]
}

Rules:
- Each requirement becomes one test step
- Maintain order of requirements
- Treat sub-parts (1.1, 1.2) as separate steps
- Use clear, professional language for FDA/GxP compliance
- Output valid JSON only, no markdown or commentary"""


def generate_test_steps_with_claude(urs_text):
    """Use Claude API to parse URS and generate test steps"""
    master_prompt = load_master_prompt()

    # Construct the prompt
    prompt = f"""{master_prompt}

URS Text:
{urs_text}

Generate the JSON output with test_steps array now:"""

    # Call Claude API
    message = client.messages.create(
        model="claude-sonnet-4-5-20250929",
        max_tokens=8000,
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    # Extract the response
    response_text = message.content[0].text

    # Try to parse JSON from the response
    # Sometimes Claude might wrap it in markdown code blocks
    response_text = response_text.strip()

    # Remove markdown code blocks if present
    if response_text.startswith('```json'):
        response_text = response_text[7:]
    if response_text.startswith('```'):
        response_text = response_text[3:]
    if response_text.endswith('```'):
        response_text = response_text[:-3]

    response_text = response_text.strip()

    # Parse JSON
    try:
        test_data = json.loads(response_text)
        return test_data
    except json.JSONDecodeError as e:
        print(f"JSON parse error: {e}")
        print(f"Response: {response_text}")
        raise ValueError(f"Failed to parse JSON from Claude response: {e}")


def populate_word_template(template_path, test_steps):
    """Populate the Word template with test steps"""
    doc = Document(template_path)

    # Find the table in the document (assumes first table is the test script table)
    if len(doc.tables) == 0:
        raise ValueError("No table found in template document")

    table = doc.tables[0]

    # Verify table has correct headers (at least requirement columns)
    # Expected columns: Step | Requirement # | Description | Expected Result | Pass/Fail | Initial | Date

    # Add rows for each test step
    for step in test_steps:
        row_cells = table.add_row().cells

        # Populate cells (adjust indices based on your template structure)
        if len(row_cells) >= 7:
            row_cells[0].text = str(step.get('step_no', ''))
            row_cells[1].text = step.get('requirement_id', '')
            row_cells[2].text = step.get('description', '')
            row_cells[3].text = step.get('expected_result', '')
            # Leave Pass/Fail, Initial, Date empty for manual completion
            row_cells[4].text = ''
            row_cells[5].text = ''
            row_cells[6].text = ''
        elif len(row_cells) >= 4:
            # Minimal template
            row_cells[0].text = str(step.get('step_no', ''))
            row_cells[1].text = step.get('requirement_id', '')
            row_cells[2].text = step.get('description', '')
            row_cells[3].text = step.get('expected_result', '')

    return doc


@app.route('/')
def landing():
    """Landing page - redirects logged-in users to dashboard"""
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return render_template('landing.html')


@app.route('/dashboard')
@login_required
def dashboard():
    """User dashboard with project list"""
    projects = Project.query.filter_by(user_id=current_user.id).order_by(Project.updated_at.desc()).all()

    # Calculate stats
    stats = {
        'total_projects': len(projects),
        'completed_projects': sum(1 for p in projects if p.status == 'executed'),
        'draft_projects': sum(1 for p in projects if p.status in ['draft', 'in_review']),
        'time_saved': len(projects) * 400  # Estimate 400 hours saved per project
    }

    return render_template('dashboard.html', projects=projects, stats=stats)


@app.route('/generator')
@login_required
def generator():
    """Test script generator page"""
    project_id = request.args.get('project_id')
    project = None

    if project_id:
        project = Project.query.filter_by(project_id=project_id, user_id=current_user.id).first()

    return render_template('generator.html', project=project)


@app.route('/projects/<project_id>')
@login_required
def view_project(project_id):
    """View a specific project"""
    project = Project.query.filter_by(project_id=project_id, user_id=current_user.id).first_or_404()

    # Parse test steps from JSON
    test_steps = []
    if project.test_steps:
        try:
            test_steps = json.loads(project.test_steps)
        except:
            test_steps = []

    # Get version history
    versions = project.versions.all()

    return render_template('project_view.html', project=project, test_steps=test_steps, versions=versions)


@app.route('/api/projects/<project_id>', methods=['DELETE'])
@login_required
def delete_project(project_id):
    """Delete a project"""
    try:
        project = Project.query.filter_by(project_id=project_id, user_id=current_user.id).first_or_404()
        db.session.delete(project)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/generate-preview', methods=['POST'])
def generate_preview():
    """Generate test steps and return for preview/editing"""
    try:
        # Check if files are present
        if 'urs_file' not in request.files:
            return jsonify({'error': 'No URS file uploaded'}), 400

        if 'template_file' not in request.files:
            return jsonify({'error': 'No template file uploaded'}), 400

        urs_file = request.files['urs_file']
        template_file = request.files['template_file']

        # Validate files
        if urs_file.filename == '':
            return jsonify({'error': 'No URS file selected'}), 400

        if template_file.filename == '':
            return jsonify({'error': 'No template file selected'}), 400

        if not allowed_file(urs_file.filename, ALLOWED_URS_EXTENSIONS):
            return jsonify({'error': 'Invalid URS file type. Allowed: PDF, DOCX, TXT'}), 400

        if not allowed_file(template_file.filename, ALLOWED_TEMPLATE_EXTENSIONS):
            return jsonify({'error': 'Invalid template file type. Allowed: DOCX'}), 400

        # Save files temporarily with unique names
        session_id = str(uuid.uuid4())
        urs_filename = f"{session_id}_{secure_filename(urs_file.filename)}"
        template_filename = f"{session_id}_{secure_filename(template_file.filename)}"

        urs_path = os.path.join(app.config['UPLOAD_FOLDER'], urs_filename)
        template_path = os.path.join(app.config['UPLOAD_FOLDER'], template_filename)

        urs_file.save(urs_path)
        template_file.save(template_path)

        print(f"Saved URS to: {urs_path}")
        print(f"Saved template to: {template_path}")

        # Step 1: Extract text from URS
        print("Extracting text from URS...")
        urs_text = extract_urs_text(urs_path, urs_file.filename)

        if not urs_text or not urs_text.strip():
            # Clean up files
            try:
                os.remove(urs_path)
                os.remove(template_path)
            except:
                pass
            return jsonify({'error': 'No text could be extracted from URS document. Please ensure it is a valid text-based document.'}), 400

        # Step 2: Generate test steps using Claude
        print("Generating test steps with Claude...")
        if not client:
            return jsonify({'error': 'Anthropic API key not configured. Please set ANTHROPIC_API_KEY in .env file.'}), 500

        test_data = generate_test_steps_with_claude(urs_text)

        if 'test_steps' not in test_data:
            return jsonify({'error': 'Invalid response from AI: missing test_steps'}), 500

        # Store template path and URS text in session for later download
        session['template_path'] = template_path
        session['session_id'] = session_id
        session['urs_text'] = urs_text  # Store for enhanced features

        # Save project to database if user is authenticated
        if current_user.is_authenticated:
            try:
                # Create or update project
                project = Project(
                    project_id=session_id,
                    name=f"Project - {urs_file.filename}",
                    description=f"Generated from {urs_file.filename}",
                    user_id=current_user.id,
                    urs_filename=urs_file.filename,
                    urs_text=urs_text,
                    template_filename=template_file.filename,
                    test_steps=json.dumps(test_data['test_steps']),
                    status='draft'
                )
                db.session.add(project)
                db.session.commit()
                print(f"✓ Project saved to database: {project.project_id}")
            except Exception as e:
                print(f"⚠ Failed to save project to database: {e}")
                # Don't fail the request if database save fails

        # Clean up URS file (we don't need it anymore in uploads folder)
        try:
            os.remove(urs_path)
        except:
            pass

        # Return test steps for preview/editing
        return jsonify({
            'test_steps': test_data['test_steps'],
            'session_id': session_id
        })

    except Exception as e:
        print(f"Error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': f'Server error: {str(e)}'}), 500


@app.route('/api/generate-final', methods=['POST'])
def generate_final():
    """Generate final Word document from edited test steps"""
    try:
        data = request.get_json()

        if not data or 'test_steps' not in data:
            return jsonify({'error': 'No test steps provided'}), 400

        test_steps = data['test_steps']
        session_id = data.get('session_id')

        # Get template path from session or use default
        template_path = session.get('template_path')

        if not template_path or not os.path.exists(template_path):
            return jsonify({'error': 'Template file not found. Please upload files again.'}), 400

        # Populate Word template
        print("Populating Word template...")
        populated_doc = populate_word_template(template_path, test_steps)

        # Save to BytesIO
        output = BytesIO()
        populated_doc.save(output)
        output.seek(0)

        # Clean up template file
        try:
            os.remove(template_path)
            session.pop('template_path', None)
            session.pop('session_id', None)
        except:
            pass

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='generated_test_script.docx'
        )

    except Exception as e:
        print(f"Error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': f'Server error: {str(e)}'}), 500


@app.route('/api/preview', methods=['POST'])
def preview_test_steps():
    """Preview test steps without generating Word doc"""
    try:
        if 'urs_file' not in request.files:
            return jsonify({'error': 'No URS file uploaded'}), 400

        urs_file = request.files['urs_file']

        if urs_file.filename == '':
            return jsonify({'error': 'No URS file selected'}), 400

        if not allowed_file(urs_file.filename, ALLOWED_URS_EXTENSIONS):
            return jsonify({'error': 'Invalid URS file type. Allowed: PDF, DOCX, TXT'}), 400

        # Save file temporarily
        urs_filename = secure_filename(urs_file.filename)
        urs_path = os.path.join(app.config['UPLOAD_FOLDER'], urs_filename)
        urs_file.save(urs_path)

        # Extract text from URS
        urs_text = extract_urs_text(urs_path, urs_filename)

        if not urs_text.strip():
            return jsonify({'error': 'No text could be extracted from URS document'}), 400

        # Generate test steps using Claude
        test_data = generate_test_steps_with_claude(urs_text)

        # Clean up
        try:
            os.remove(urs_path)
        except:
            pass

        return jsonify(test_data)

    except Exception as e:
        print(f"Error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': f'Server error: {str(e)}'}), 500


# ═══════════════════════════════════════════════════════════════════════════════
# ENHANCED VALIDATION FEATURES (New Routes - Do NOT modify existing routes above)
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/api/generate-rtm-excel', methods=['POST'])
def generate_rtm_excel_endpoint():
    """Generate RTM as Excel file"""
    if not ENHANCED_FEATURES_AVAILABLE:
        return jsonify({'error': 'Enhanced features not available'}), 503

    try:
        data = request.get_json()
        test_steps = data.get('test_steps', [])

        if not test_steps:
            return jsonify({'error': 'No test steps provided'}), 400

        excel_file = generate_rtm_excel(test_steps)

        return send_file(
            excel_file,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='RTM_Requirements_Traceability_Matrix.xlsx'
        )

    except Exception as e:
        print(f"RTM Excel Error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': f'Failed to generate RTM: {str(e)}'}), 500


@app.route('/api/generate-rtm-word', methods=['POST'])
def generate_rtm_word_endpoint():
    """Generate RTM as Word document"""
    if not ENHANCED_FEATURES_AVAILABLE:
        return jsonify({'error': 'Enhanced features not available'}), 503

    try:
        data = request.get_json()
        test_steps = data.get('test_steps', [])

        if not test_steps:
            return jsonify({'error': 'No test steps provided'}), 400

        word_file = generate_rtm_word(test_steps)

        return send_file(
            word_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='RTM_Requirements_Traceability_Matrix.docx'
        )

    except Exception as e:
        print(f"RTM Word Error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': f'Failed to generate RTM: {str(e)}'}), 500


@app.route('/api/analyze-changes', methods=['POST'])
def analyze_changes_endpoint():
    """Analyze changes between URS versions"""
    if not ENHANCED_FEATURES_AVAILABLE:
        return jsonify({'error': 'Enhanced features not available'}), 503

    try:
        data = request.get_json()
        urs_text = data.get('urs_text', '')
        test_steps = data.get('test_steps', [])
        project_name = data.get('project_name', 'default')
        save_baseline = data.get('save_baseline', True)

        if not urs_text or not test_steps:
            return jsonify({'error': 'URS text and test steps required'}), 400

        # Analyze changes
        change_report = change_analyzer.analyze_changes(urs_text, test_steps, project_name)

        # Save as baseline if requested
        if save_baseline:
            baseline_id = change_analyzer.save_baseline(urs_text, test_steps, project_name)
            change_report['new_baseline_id'] = baseline_id

        return jsonify(change_report)

    except Exception as e:
        print(f"Change Analysis Error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': f'Failed to analyze changes: {str(e)}'}), 500


@app.route('/api/generate-validation-plan', methods=['POST'])
def generate_validation_plan_endpoint():
    """Generate Validation Master Plan"""
    if not ENHANCED_FEATURES_AVAILABLE:
        return jsonify({'error': 'Enhanced features not available'}), 503

    try:
        data = request.get_json()
        test_steps = data.get('test_steps', [])
        session_id = data.get('session_id', '')

        # Try to get URS text from session
        urs_text = session.get('urs_text', '')

        # If not in session and user is authenticated, try to get from database
        if not urs_text and current_user.is_authenticated and session_id:
            project = Project.query.filter_by(project_id=session_id, user_id=current_user.id).first()
            if project:
                urs_text = project.urs_text

        if not urs_text or not test_steps:
            return jsonify({'error': 'URS text and test steps required'}), 400

        vmp_file = validation_doc_gen.generate_validation_plan(urs_text, test_steps)

        return send_file(
            vmp_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='Validation_Master_Plan.docx'
        )

    except Exception as e:
        print(f"VMP Generation Error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': f'Failed to generate validation plan: {str(e)}'}), 500


@app.route('/api/generate-validation-summary', methods=['POST'])
def generate_validation_summary_endpoint():
    """Generate Validation Summary Report"""
    if not ENHANCED_FEATURES_AVAILABLE:
        return jsonify({'error': 'Enhanced features not available'}), 503

    try:
        data = request.get_json()
        test_steps = data.get('test_steps', [])

        if not test_steps:
            return jsonify({'error': 'Test steps required'}), 400

        vsr_file = validation_doc_gen.generate_validation_summary(test_steps)

        return send_file(
            vsr_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='Validation_Summary_Report.docx'
        )

    except Exception as e:
        print(f"VSR Generation Error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': f'Failed to generate validation summary: {str(e)}'}), 500


@app.route('/api/check-quality', methods=['POST'])
def check_quality_endpoint():
    """Check requirements quality"""
    if not ENHANCED_FEATURES_AVAILABLE:
        return jsonify({'error': 'Enhanced features not available'}), 503

    try:
        data = request.get_json()
        session_id = data.get('session_id', '')

        # Try to get URS text from session
        urs_text = session.get('urs_text', '')

        # If not in session and user is authenticated, try to get from database
        if not urs_text and current_user.is_authenticated and session_id:
            project = Project.query.filter_by(project_id=session_id, user_id=current_user.id).first()
            if project:
                urs_text = project.urs_text

        if not urs_text:
            return jsonify({'error': 'URS text not found. Please generate a test script first.'}), 400

        quality_report = quality_checker.analyze_requirements(urs_text)

        return jsonify(quality_report)

    except Exception as e:
        print(f"Quality Check Error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': f'Failed to check quality: {str(e)}'}), 500


@app.route('/api/export-audit-package', methods=['POST'])
def export_audit_package_endpoint():
    """Export complete audit package as ZIP"""
    if not ENHANCED_FEATURES_AVAILABLE:
        return jsonify({'error': 'Enhanced features not available'}), 503

    try:
        data = request.get_json()

        # Required data
        test_steps = data.get('test_steps', [])
        urs_text = data.get('urs_text', '')

        if not test_steps or not urs_text:
            return jsonify({'error': 'Test steps and URS text required'}), 400

        # Generate all components
        package_data = {
            'test_steps': test_steps,
            'urs_text': urs_text
        }

        # Generate RTM files
        try:
            package_data['rtm_excel'] = generate_rtm_excel(test_steps)
            package_data['rtm_word'] = generate_rtm_word(test_steps)
        except Exception as e:
            print(f"Warning: RTM generation failed: {e}")

        # Generate validation documents
        try:
            package_data['validation_plan'] = validation_doc_gen.generate_validation_plan(urs_text, test_steps)
            package_data['validation_summary'] = validation_doc_gen.generate_validation_summary(test_steps)
        except Exception as e:
            print(f"Warning: Validation docs generation failed: {e}")

        # Generate quality report
        try:
            package_data['quality_report'] = quality_checker.analyze_requirements(urs_text)
        except Exception as e:
            print(f"Warning: Quality check failed: {e}")

        # Create audit package ZIP
        zip_file = audit_exporter.create_audit_package(**package_data)

        return send_file(
            zip_file,
            mimetype='application/zip',
            as_attachment=True,
            download_name='Audit_Package_Complete.zip'
        )

    except Exception as e:
        print(f"Audit Package Error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': f'Failed to create audit package: {str(e)}'}), 500


# ═══════════════════════════════════════════════════════════════════════════════
# END OF ENHANCED FEATURES
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
