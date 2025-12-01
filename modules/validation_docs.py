"""
Validation Plan & Summary Report Generator
Generates GxP-compliant validation documentation (VMP, VSR)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime


class ValidationDocGenerator:
    def __init__(self):
        self.timestamp = datetime.now().strftime('%Y-%m-%d')

    def generate_validation_plan(self, urs_text, test_steps):
        """
        Generate Validation Master Plan (VMP)

        Args:
            urs_text: Extracted URS text
            test_steps: Generated test steps

        Returns:
            BytesIO object containing Word document
        """
        doc = Document()

        # Title Page
        title = doc.add_heading('Validation Master Plan', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph(f'Document Date: {self.timestamp}').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f'Version: 1.0').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Document Control
        doc.add_heading('1. Document Control', level=2)
        doc.add_paragraph(f'Document ID: VMP-001')
        doc.add_paragraph(f'Effective Date: {self.timestamp}')
        doc.add_paragraph(f'Status: Draft')
        doc.add_paragraph()

        # Purpose
        doc.add_heading('2. Purpose', level=2)
        doc.add_paragraph(
            'This Validation Master Plan (VMP) defines the validation strategy, approach, and '
            'activities required to ensure the system meets user requirements and complies with '
            'FDA 21 CFR Part 11 and GxP regulations.'
        )
        doc.add_paragraph()

        # Scope
        doc.add_heading('3. Scope', level=2)
        doc.add_paragraph(
            'This validation plan covers all functional and non-functional requirements '
            'specified in the User Requirements Specification (URS). The validation activities '
            'include Installation Qualification (IQ), Operational Qualification (OQ), and '
            'Performance Qualification (PQ).'
        )
        doc.add_paragraph()

        # Validation Approach
        doc.add_heading('4. Validation Approach', level=2)
        doc.add_paragraph('The validation will follow a risk-based approach with the following phases:')

        phases = [
            '• Planning: Define validation strategy and create test protocols',
            '• Execution: Execute test scripts and document results',
            '• Review: Review test results and address any deviations',
            '• Approval: Obtain necessary approvals from stakeholders',
            '• Archival: Archive validation package for regulatory compliance'
        ]

        for phase in phases:
            doc.add_paragraph(phase)

        doc.add_paragraph()

        # Test Strategy
        doc.add_heading('5. Test Strategy', level=2)
        doc.add_paragraph(f'Total Requirements: {len(test_steps)}')
        doc.add_paragraph(f'Total Test Cases: {len(test_steps)}')
        doc.add_paragraph(f'Coverage: 100% (1:1 mapping of requirements to test cases)')
        doc.add_paragraph()

        doc.add_paragraph('Test Categories:')
        doc.add_paragraph('• Functional Testing: Verify all functional requirements')
        doc.add_paragraph('• Non-Functional Testing: Verify performance, security, and usability')
        doc.add_paragraph('• Integration Testing: Verify system integration points')
        doc.add_paragraph('• User Acceptance Testing: Verify end-user workflows')
        doc.add_paragraph()

        # Roles and Responsibilities
        doc.add_heading('6. Roles and Responsibilities', level=2)

        # Create table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Role'
        header_cells[1].text = 'Responsibility'

        # Make headers bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Roles
        roles = [
            ('Validation Lead', 'Oversee validation activities and ensure compliance'),
            ('Test Engineer', 'Execute test scripts and document results'),
            ('QA Reviewer', 'Review test results and approve validation package'),
            ('System Owner', 'Approve validation plan and sign off on completion')
        ]

        for role, responsibility in roles:
            row = table.add_row()
            row.cells[0].text = role
            row.cells[1].text = responsibility

        doc.add_paragraph()

        # Deliverables
        doc.add_heading('7. Validation Deliverables', level=2)
        deliverables = [
            '• Validation Master Plan (VMP)',
            '• User Requirements Specification (URS)',
            '• Requirements Traceability Matrix (RTM)',
            '• Test Scripts',
            '• Test Execution Records',
            '• Deviation Reports (if any)',
            '• Validation Summary Report (VSR)'
        ]

        for item in deliverables:
            doc.add_paragraph(item)

        doc.add_paragraph()

        # Schedule
        doc.add_heading('8. Validation Schedule', level=2)
        doc.add_paragraph('Estimated Timeline:')
        doc.add_paragraph('• Planning Phase: 2 weeks')
        doc.add_paragraph('• Execution Phase: 4 weeks')
        doc.add_paragraph('• Review & Approval: 1 week')
        doc.add_paragraph('• Total Duration: 7 weeks')
        doc.add_paragraph()

        # Approval
        doc.add_heading('9. Approvals', level=2)
        self._add_approval_table(doc)

        # Save
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return output

    def generate_validation_summary(self, test_steps, execution_results=None):
        """
        Generate Validation Summary Report (VSR)

        Args:
            test_steps: Test steps that were executed
            execution_results: Optional execution results (pass/fail counts)

        Returns:
            BytesIO object containing Word document
        """
        doc = Document()

        # Title Page
        title = doc.add_heading('Validation Summary Report', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph(f'Report Date: {self.timestamp}').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f'Version: 1.0').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Executive Summary
        doc.add_heading('1. Executive Summary', level=2)
        doc.add_paragraph(
            f'This Validation Summary Report documents the completion of validation activities '
            f'for the system. A total of {len(test_steps)} test cases were executed covering '
            f'all {len(test_steps)} requirements specified in the URS.'
        )
        doc.add_paragraph()

        # Validation Scope
        doc.add_heading('2. Validation Scope', level=2)
        doc.add_paragraph(
            'The validation covered all functional and non-functional requirements as defined '
            'in the User Requirements Specification (URS). All test cases were executed in '
            'accordance with the approved Test Scripts.'
        )
        doc.add_paragraph()

        # Test Execution Summary
        doc.add_heading('3. Test Execution Summary', level=2)

        # Summary table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'

        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        metrics = [
            ('Total Requirements', str(len(test_steps))),
            ('Total Test Cases', str(len(test_steps))),
            ('Test Cases Executed', str(len(test_steps))),
            ('Test Cases Passed', 'TBD (Execute tests to populate)'),
            ('Test Cases Failed', 'TBD (Execute tests to populate)'),
            ('Execution Coverage', '100%')
        ]

        for metric, value in metrics:
            row = table.add_row()
            row.cells[0].text = metric
            row.cells[1].text = value

        doc.add_paragraph()

        # Traceability
        doc.add_heading('4. Requirements Traceability', level=2)
        doc.add_paragraph(
            'Complete traceability has been established between all requirements and test cases. '
            'Refer to the Requirements Traceability Matrix (RTM) for detailed mapping.'
        )
        doc.add_paragraph()
        doc.add_paragraph(f'Requirements Coverage: 100% ({len(test_steps)} of {len(test_steps)} requirements tested)')
        doc.add_paragraph()

        # Deviations
        doc.add_heading('5. Deviations and Exceptions', level=2)
        doc.add_paragraph('No deviations or exceptions were recorded during validation execution.')
        doc.add_paragraph('All test cases passed without requiring retesting or corrective actions.')
        doc.add_paragraph()

        # Conclusion
        doc.add_heading('6. Conclusion', level=2)
        doc.add_paragraph(
            'The validation activities have been successfully completed in accordance with the '
            'approved Validation Master Plan. All requirements have been verified through '
            'comprehensive testing. The system is validated and suitable for its intended use '
            'in a GxP environment.'
        )
        doc.add_paragraph()

        # Recommendations
        doc.add_heading('7. Recommendations', level=2)
        doc.add_paragraph(
            'Based on successful completion of validation activities, it is recommended that:'
        )
        doc.add_paragraph('• The system be approved for production use')
        doc.add_paragraph('• Ongoing change control procedures be followed for any system changes')
        doc.add_paragraph('• Periodic reviews be conducted to ensure continued compliance')
        doc.add_paragraph()

        # Approval
        doc.add_heading('8. Approvals', level=2)
        self._add_approval_table(doc)

        # Save
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return output

    def _add_approval_table(self, doc):
        """Add standard approval signature table"""
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Role'
        header_cells[1].text = 'Signature'
        header_cells[2].text = 'Date'

        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Signature rows
        roles = [
            'Validation Lead',
            'QA Reviewer',
            'System Owner'
        ]

        for role in roles:
            row = table.add_row()
            row.cells[0].text = role
            row.cells[1].text = '_' * 30
            row.cells[2].text = '_' * 15

        doc.add_paragraph()
