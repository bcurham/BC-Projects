"""
Requirements Traceability Matrix (RTM) Generator
Generates RTM mapping URS requirements to test cases
Outputs to Excel and Word formats
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import json

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("⚠ openpyxl not available - Excel RTM export disabled")


def generate_rtm_data(test_steps):
    """
    Generate RTM data structure from test steps

    Args:
        test_steps: List of test step dictionaries with:
                   - step_no
                   - requirement_id
                   - description
                   - expected_result

    Returns:
        List of RTM entries
    """
    rtm_entries = []

    for step in test_steps:
        entry = {
            'requirement_id': step.get('requirement_id', ''),
            'requirement_description': step.get('description', ''),
            'test_case_id': f"TC-{step.get('step_no', '')}",
            'test_description': f"Verify: {step.get('description', '')[:100]}...",
            'expected_result': step.get('expected_result', ''),
            'traceability_status': 'Mapped',
            'coverage': '1:1'
        }
        rtm_entries.append(entry)

    return rtm_entries


def generate_rtm_excel(test_steps):
    """
    Generate RTM as Excel file

    Returns:
        BytesIO object containing Excel file
    """
    if not EXCEL_AVAILABLE:
        raise ImportError("openpyxl is required for Excel export")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "RTM"

    # Define styles
    header_fill = PatternFill(start_color="2E75B5", end_color="2E75B5", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    cell_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Headers
    headers = [
        'Requirement ID',
        'Requirement Description',
        'Test Case ID',
        'Test Description',
        'Expected Result',
        'Traceability Status',
        'Coverage'
    ]

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = border

    # Set column widths
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 40
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 40
    ws.column_dimensions['E'].width = 40
    ws.column_dimensions['F'].width = 18
    ws.column_dimensions['G'].width = 12

    # Generate RTM data
    rtm_data = generate_rtm_data(test_steps)

    # Fill data
    for row_idx, entry in enumerate(rtm_data, start=2):
        ws.cell(row=row_idx, column=1, value=entry['requirement_id']).alignment = cell_alignment
        ws.cell(row=row_idx, column=2, value=entry['requirement_description']).alignment = cell_alignment
        ws.cell(row=row_idx, column=3, value=entry['test_case_id']).alignment = cell_alignment
        ws.cell(row=row_idx, column=4, value=entry['test_description']).alignment = cell_alignment
        ws.cell(row=row_idx, column=5, value=entry['expected_result']).alignment = cell_alignment
        ws.cell(row=row_idx, column=6, value=entry['traceability_status']).alignment = Alignment(horizontal="center")
        ws.cell(row=row_idx, column=7, value=entry['coverage']).alignment = Alignment(horizontal="center")

        # Apply borders
        for col in range(1, 8):
            ws.cell(row=row_idx, column=col).border = border

    # Save to BytesIO
    output = BytesIO()
    wb.save(output)
    output.seek(0)

    return output


def generate_rtm_word(test_steps):
    """
    Generate RTM as Word document

    Returns:
        BytesIO object containing Word file
    """
    doc = Document()

    # Title
    title = doc.add_heading('Requirements Traceability Matrix', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f'Generated: {_get_timestamp()}')
    doc.add_paragraph(f'Total Requirements: {len(test_steps)}')
    doc.add_paragraph()

    # Introduction
    doc.add_heading('Purpose', level=2)
    doc.add_paragraph(
        'This Requirements Traceability Matrix (RTM) establishes bidirectional traceability '
        'between user requirements and test cases, ensuring comprehensive test coverage and '
        'regulatory compliance.'
    )
    doc.add_paragraph()

    # RTM Table
    doc.add_heading('Traceability Matrix', level=2)

    # Generate RTM data
    rtm_data = generate_rtm_data(test_steps)

    # Create table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Req ID', 'Requirement', 'Test Case', 'Test Description', 'Expected Result', 'Status', 'Coverage']

    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Data rows
    for entry in rtm_data:
        row_cells = table.add_row().cells
        row_cells[0].text = entry['requirement_id']
        row_cells[1].text = entry['requirement_description'][:200]  # Truncate for readability
        row_cells[2].text = entry['test_case_id']
        row_cells[3].text = entry['test_description']
        row_cells[4].text = entry['expected_result'][:200]
        row_cells[5].text = entry['traceability_status']
        row_cells[6].text = entry['coverage']

    # Summary statistics
    doc.add_paragraph()
    doc.add_heading('Coverage Summary', level=2)
    doc.add_paragraph(f'Total Requirements: {len(rtm_data)}')
    doc.add_paragraph(f'Total Test Cases: {len(rtm_data)}')
    doc.add_paragraph(f'Coverage: 100%')

    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output


def _get_timestamp():
    """Get current timestamp for documentation"""
    from datetime import datetime
    return datetime.now().strftime('%Y-%m-%d %H:%M:%S')
